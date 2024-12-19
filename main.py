# import os
# import re
# import json
# import tempfile
# import logging
# from flask import Flask, request, render_template, flash, send_file, url_for
# from docx import Document as DocxDocument
# from selenium_utils import init_driver, process_bot
# import time

# # # Logging Configuration
# # logging.basicConfig(
# #     filename='app.log',
# #     level=logging.INFO,
# #     format='%(asctime)s %(levelname)s %(message)s',
# #     datefmt='%Y-%m-%d %H:%M:%S'
# # )

# app = Flask(__name__)
# app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')  # Replace with a secure secret key
# app.config['UPLOAD_FOLDER'] = 'uploads'

# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'config.json')

# if not os.path.exists(CONFIG_FILE):
#     logging.critical(f"Config file {CONFIG_FILE} not found.")
#     raise FileNotFoundError(f"Config file {CONFIG_FILE} not found.")

# try:
#     with open(CONFIG_FILE, 'r') as f:
#         GPT_BOTS = json.load(f)
#     logging.info("Loaded GPT_BOTS configuration successfully.")
# except json.JSONDecodeError as e:
#     logging.critical(f"Error parsing GPT_BOTS JSON: {e}")
#     raise Exception(f"Error parsing GPT_BOTS JSON: {e}")

# @app.route('/', methods=['GET', 'POST'])
# def process_text():
#     matched_codes = []
#     combined_text = ""
#     prompts = [
#         {
#             "prompt": "Create an outline for the assignment. Provide a breakdown of sections ...",
#             "expected": "A detailed outline..."
#         },
#         {
#             "prompt": "Draft the detailed sections till the first 100 words ...",
#             "expected": "A well-rounded introduction..."
#         },
#         {
#             "prompt": "Kindly provide the complete list of references used in the above 100 words ...",
#             "expected": "A well-formatted list of references."
#         },
#         {
#             "prompt": "Draft the detailed next sections till the 2nd 100 words ...",
#             "expected": "A well-rounded next section..."
#         },
#         {
#             "prompt": "Draft the detailed next sections till the 3rd 100 words ...",
#             "expected": "A well-rounded next section..."
#         },
#         {
#             "prompt": "Draft the detailed next sections till the conclusion ...",
#             "expected": "A well-rounded next section till Conclusion..."
#         }
#     ]

#     if request.method == 'POST':
#         user_text = request.form.get('user_text', '').strip()
#         if not user_text:
#             flash('No text entered. Please provide input.')
#             logging.warning("User submitted empty text input.")
#             return render_template('input.html')

#         # Check for matches in the user text
#         found_codes = []
#         for code in GPT_BOTS.keys():
#             if re.search(rf'\b{re.escape(code)}\b', user_text, re.IGNORECASE):
#                 found_codes.append(code)

#         if not found_codes:
#             flash('No matching code found in the provided text.')
#             logging.info("No matching codes found in user input.")
#             return render_template('input.html')

#         logging.info(f"Found matching codes: {found_codes}")

#         all_responses = []
#         # Process each matched code
#         for code in found_codes:
#             bot_url = GPT_BOTS[code]
#             logging.info(f"Processing bot for code: {code} with URL: {bot_url}")

#             driver = init_driver()
#             try:
#                 time.sleep(30)
#                 bot_responses = process_bot(driver, bot_url, prompts)
#                 all_responses.extend(bot_responses)
#                 logging.info(f"Collected {len(bot_responses)} responses from bot {code}.")
#             except Exception as e:
#                 logging.error(f"Error processing bot for code {code}: {e}")
#                 flash(f"Error processing bot for code {code}: {e}")
#             finally:
#                 try:
#                     driver.quit()
#                     logging.info(f"Closed Selenium WebDriver for bot {code}.")
#                 except Exception as e:
#                     logging.error(f"Error closing WebDriver for bot {code}: {e}")

#             matched_codes.append({"code": code, "bot_url": bot_url})

#         if not all_responses:
#             flash("No responses were generated from the bots.")
#             logging.warning("No responses were collected from any bot.")
#             return render_template('input.html', matched_codes=matched_codes)

#         # Combine all responses

#         for response in all_responses:
#             combined_text += f"\n\n**Prompt:** {response['prompt']}\n\n**Response:** {response['answer']}"

#         # Create a Word document
#         try:
#             doc = DocxDocument()
#             doc.add_heading('Bot Responses', 0)
#             doc.add_paragraph(combined_text)

#             output_filename = "Bot_Responses.docx"
#             output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#             doc.save(output_path)

#             flash('Processing complete. Download the generated Word document below.')
#             logging.info("Word document created successfully.")
#             return render_template('input.html', 
#                                    matched_codes=matched_codes,
#                                    combined_text=combined_text,
#                                    download_link=output_filename)
#         except Exception as e:
#             logging.error(f"Error creating Word document: {e}")
#             flash(f"Error creating Word document: {e}")
#             return render_template('input.html', matched_codes=matched_codes, combined_text=combined_text)

#     # GET request
#     return render_template('input.html')

# @app.route('/download/<filename>')
# def download_file(filename):
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     if os.path.exists(file_path):
#         logging.info(f"User requested download for file: {filename}")
#         return send_file(file_path, as_attachment=True, download_name=filename)
#     else:
#         flash('File not found.')
#         logging.warning(f"Download attempted for non-existent file: {filename}")
#         return render_template('input.html')

# if __name__ == '__main__':
#     app.run(debug=True)


# import os
# import re
# import json
# import tempfile
# import logging
# from flask import Flask, request, render_template, flash, send_file, url_for
# from docx import Document as DocxDocument
# from selenium_utils import init_driver, process_bot
# app = Flask(__name__)
# app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')  # Replace with a secure secret key
# app.config['UPLOAD_FOLDER'] = 'uploads'
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
# CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'config.json')

# if not os.path.exists(CONFIG_FILE):
#     logging.critical(f"Config file {CONFIG_FILE} not found.")
#     raise FileNotFoundError(f"Config file {CONFIG_FILE} not found.")

# try:
#     with open(CONFIG_FILE, 'r') as f:
#         GPT_BOTS = json.load(f)
#     logging.info("Loaded GPT_BOTS configuration successfully.")
# except json.JSONDecodeError as e:
#     logging.critical(f"Error parsing GPT_BOTS JSON: {e}")
#     raise Exception(f"Error parsing GPT_BOTS JSON: {e}")

# @app.route('/', methods=['GET', 'POST'])
# def process_text():
#     matched_codes = []
#     combined_text = ""
#     prompts = [
#             {
#                     "prompt": "Create an outline for the assignment. Provide a breakdown of sections with a suggested word count for each, totalling to required words as per the Assignment Brief. Ensure that the outline aligns with main focus, requirements and includes recommended headings or subheadings.",
#                     "expected": "A detailed outline with sections, word allocations, and subheadings tailored to the topic."
#             },
#             {
#                     "prompt": "Draft the detailed sections till the first 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief.",
#                     "expected": "A well-rounded introduction with relevant background, 2-3 in-text citations, and properly formatted references."
#             },
#             {
#                     "prompt": "Kindly provide the complete list of references used in the above 100 words in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the 2nd 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 2nd 1000 words in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the 3rd 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 3rd 1000 words in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the conclusion of the Assignment in the allocated word count according to the outline, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide the complete list of references used in these sections in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s till Conclusion, 2-3 in-text citations, and properly formatted references and well-formatted list of references."
#             }
#         ]

#     if request.method == 'POST':
#         user_text = request.form.get('user_text', '').strip()
#         if not user_text:
#             flash('No text entered. Please provide input.')
#             logging.warning("User submitted empty text input.")
#             return render_template('input.html')

#         # Check for matches in the user text
#         found_codes = []
#         for code in GPT_BOTS.keys():
#             if re.search(rf'\b{re.escape(code)}\b', user_text, re.IGNORECASE):
#                 found_codes.append(code)

#         if not found_codes:
#             flash('No matching code found in the provided text.')
#             logging.info("No matching codes found in user input.")
#             return render_template('input.html')

#         logging.info(f"Found matching codes: {found_codes}")

#         all_responses = []
#         # Process each matched code
#         for code in found_codes:
#             bot_url = GPT_BOTS[code]
#             logging.info(f"Processing bot for code: {code} with URL: {bot_url}")

#             driver = init_driver()
#             try:
#                 bot_responses = process_bot(driver, bot_url, prompts)
#                 if bot_responses:
#                     all_responses.extend(bot_responses)
#                     logging.info(f"Collected {len(bot_responses)} responses from bot {code}.")
#                 else:
#                     logging.warning(f"No responses collected from bot {code}.")
#             except Exception as e:
#                 logging.error(f"Error processing bot for code {code}: {e}")
#                 flash(f"Error processing bot for code {code}: {e}")
#             finally:
#                 try:
#                     driver.quit()
#                     logging.info(f"Closed Selenium WebDriver for bot {code}.")
#                 except Exception as e:
#                     logging.error(f"Error closing WebDriver for bot {code}: {e}")

#             matched_codes.append({"code": code, "bot_url": bot_url})

#         if not all_responses:
#             flash("No responses were generated from the bots.")
#             logging.warning("No responses were collected from any bot.")
#             return render_template('input.html', matched_codes=matched_codes)

#         # Combine all responses
#         for response in all_responses:
#             combined_text += f"\n\n**Prompt:** {response['prompt']}\n\n**Response:** {response['answer']}"

#         # Create a Word document
#         try:
#             doc = DocxDocument()
#             doc.add_heading('Bot Responses', 0)
#             doc.add_paragraph(combined_text)

#             output_filename = "Bot_Responses.docx"
#             output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#             doc.save(output_path)

#             flash('Processing complete. Download the generated Word document below.')
#             logging.info("Word document created successfully.")
#             return render_template('input.html', 
#                                    matched_codes=matched_codes,
#                                    combined_text=combined_text,
#                                    download_link=output_filename)
#         except Exception as e:
#             logging.error(f"Error creating Word document: {e}")
#             flash(f"Error creating Word document: {e}")
#             return render_template('input.html', matched_codes=matched_codes, combined_text=combined_text)

#     # GET request
#     return render_template('input.html')

# @app.route('/download/<filename>')
# def download_file(filename):
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     if os.path.exists(file_path):
#         logging.info(f"User requested download for file: {filename}")
#         return send_file(file_path, as_attachment=True, download_name=filename)
#     else:
#         flash('File not found.')
#         logging.warning(f"Download attempted for non-existent file: {filename}")
#         return render_template('input.html')

# if __name__ == '__main__':
#     app.run(debug=True)

# import os
# import re
# import json
# import tempfile
# import logging
# from flask import Flask, request, render_template, flash, send_file, url_for
# from docx import Document as DocxDocument
# from selenium_utils import init_driver, process_bot

# # # Logging Configuration with Rotating File Handler
# # from logging.handlers import RotatingFileHandler

# # handler = RotatingFileHandler('app.log', maxBytes=5*1024*1024, backupCount=5)  # 5MB per log file
# # handler.setLevel(logging.INFO)
# # # formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
# # handler.setFormatter(formatter)
# # logging.getLogger().addHandler(handler)

# app = Flask(__name__)
# app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')  # Replace with a secure secret key
# app.config['UPLOAD_FOLDER'] = 'uploads'

# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'config.json')

# if not os.path.exists(CONFIG_FILE):
#     logging.critical(f"Config file {CONFIG_FILE} not found.")
#     raise FileNotFoundError(f"Config file {CONFIG_FILE} not found.")

# try:
#     with open(CONFIG_FILE, 'r') as f:
#         GPT_BOTS = json.load(f)
#     logging.info("Loaded GPT_BOTS configuration successfully.")
# except json.JSONDecodeError as e:
#     logging.critical(f"Error parsing GPT_BOTS JSON: {e}")
#     raise Exception(f"Error parsing GPT_BOTS JSON: {e}")

# @app.route('/', methods=['GET', 'POST'])
# def process_text():
#     matched_codes = []
#     combined_text = ""
#     prompts = [
#             {
#                     "prompt": "Create an outline for the assignment. Provide a breakdown of sections with a suggested word count for each, totalling to required words as per the Assignment Brief. Ensure that the outline aligns with main focus, requirements and includes recommended headings or subheadings. Do not include any disclaimers or assistance text in the response.",
#                     "expected": "A detailed outline with sections, word allocations, and subheadings tailored to the topic."
#             },
#             {
#                     "prompt": "Draft the detailed sections till the first 1000 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief. Do not include any disclaimers or assistance text in the response.",
#                     "expected": "A well-rounded introduction with relevant background, 2-3 in-text citations, and properly formatted references."
#             },
#             {
#                     "prompt": "Kindly provide the complete list of references used in the above 1000 words in the reference format mentioned in the assignment brief. Do not include any disclaimers or assistance text in the response.",
#                     "expected": "A well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the 2nd 1000 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 2nd 1000 words in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the 3rd 1000 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 3rd 1000 words in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
#             },
#             {
#                     "prompt": "Draft the detailed next sections till the conclusion of the Assignment in the allocated word count according to the outline, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide the complete list of references used in these sections in the reference format mentioned in the assignment brief.",
#                     "expected": "A well-rounded next section/s till Conclusion, 2-3 in-text citations, and properly formatted references and well-formatted list of references."
#             }
#         ]

#     if request.method == 'POST':
#         user_text = request.form.get('user_text', '').strip()
#         if not user_text:
#             flash('No text entered. Please provide input.')
#             logging.warning("User submitted empty text input.")
#             return render_template('input.html')

#         # Check for matches in the user text
#         found_codes = []
#         for code in GPT_BOTS.keys():
#             if re.search(rf'\b{re.escape(code)}\b', user_text, re.IGNORECASE):
#                 found_codes.append(code)

#         if not found_codes:
#             flash('No matching code found in the provided text.')
#             logging.info("No matching codes found in user input.")
#             return render_template('input.html')

#         logging.info(f"Found matching codes: {found_codes}")

#         all_responses = []
#         # Process each matched code
#         for code in found_codes:
#             bot_url = GPT_BOTS[code]
#             logging.info(f"Processing bot for code: {code} with URL: {bot_url}")

#             driver = init_driver()
#             try:
#                 bot_responses = process_bot(driver, bot_url, prompts)
#                 if bot_responses:
#                     all_responses.extend(bot_responses)
#                     logging.info(f"Collected {len(bot_responses)} responses from bot {code}.")
#                 else:
#                     logging.warning(f"No responses collected from bot {code}.")
#             except Exception as e:
#                 logging.error(f"Error processing bot for code {code}: {e}")
#                 flash(f"Error processing bot for code {code}: {e}")
#             finally:
#                 try:
#                     driver.quit()
#                     logging.info(f"Closed Selenium WebDriver for bot {code}.")
#                 except Exception as e:
#                     logging.error(f"Error closing WebDriver for bot {code}: {e}")

#             matched_codes.append({"code": code, "bot_url": bot_url})

#         if not all_responses:
#             flash("No responses were generated from the bots.")
#             logging.warning("No responses were collected from any bot.")
#             return render_template('input.html', matched_codes=matched_codes)

#         # # Combine all responses
#         # for response in all_responses:
#         #     combined_text += f"\n\n**Prompt:** {response['prompt']}\n\n**Response:** {response['answer']}"
#         for response in all_responses:
#             combined_text += f"\n\n{response['answer']}"
#         # Create a Word document
#         try:
#             doc = DocxDocument()
#             doc.add_heading('Bot Responses', 0)
#             doc.add_paragraph(combined_text)

#             output_filename = "Bot_Responses.docx"
#             output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#             doc.save(output_path)

#             flash('Processing complete. Download the generated Word document below.')
#             logging.info("Word document created successfully.")
#             return render_template('input.html', 
#                                    matched_codes=matched_codes,
#                                    combined_text=combined_text,
#                                    download_link=output_filename)
#         except Exception as e:
#             logging.error(f"Error creating Word document: {e}")
#             flash(f"Error creating Word document: {e}")
#             return render_template('input.html', matched_codes=matched_codes, combined_text=combined_text)

#     # GET request
#     return render_template('input.html')

# @app.route('/download/<filename>')
# def download_file(filename):
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     if os.path.exists(file_path):
#         logging.info(f"User requested download for file: {filename}")
#         return send_file(file_path, as_attachment=True, download_name=filename)
#     else:
#         flash('File not found.')
#         logging.warning(f"Download attempted for non-existent file: {filename}")
#         return render_template('input.html')

# if __name__ == '__main__':
#     app.run(debug=True)

import os
import re
import json
import tempfile
import logging
from flask import Flask, request, render_template, flash, send_file, url_for
from docx import Document as DocxDocument
from selenium_utils import init_driver, process_bot
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# # Logging Configuration with Rotating File Handler
# from logging.handlers import RotatingFileHandler

# handler = RotatingFileHandler('app.log', maxBytes=5*1024*1024, backupCount=5)  # 5MB per log file
# handler.setLevel(logging.INFO)
# # formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
# handler.setFormatter(formatter)
# logging.getLogger().addHandler(handler)

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your_default_secret_key')  # Replace with a secure secret key
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'config.json')

if not os.path.exists(CONFIG_FILE):
    logging.critical(f"Config file {CONFIG_FILE} not found.")
    raise FileNotFoundError(f"Config file {CONFIG_FILE} not found.")

try:
    with open(CONFIG_FILE, 'r') as f:
        GPT_BOTS = json.load(f)
    logging.info("Loaded GPT_BOTS configuration successfully.")
except json.JSONDecodeError as e:
    logging.critical(f"Error parsing GPT_BOTS JSON: {e}")
    raise Exception(f"Error parsing GPT_BOTS JSON: {e}")
# Function to apply the required formatting
def apply_formatting_to_word(doc):
    # Function to set font and size
    def set_font(para, font_name='Calibri', font_size=12):
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # Iterate through paragraphs to apply formatting
    for para in doc.paragraphs:
        # Apply font and size (Calibri, size 12)
        set_font(para)

        # Set paragraph alignment to left (left-justified)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Apply line spacing and spacing before/after paragraphs
        para.paragraph_format.line_spacing = 1.15  # Line spacing at 1.15
        para.paragraph_format.space_before = 0  # Space before each paragraph
        para.paragraph_format.space_after = 10  # Space after each paragraph

    # Apply formatting for headings
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.text = para.text.upper()  # Convert text to upper case
            for run in para.runs:
                run.bold = True
                set_font(para, font_name='Arial')  # You can change to Calibri or Times New Roman

        elif para.style.name == 'Heading 2':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.text = para.text.capitalize()  # Sentence case
            for run in para.runs:
                run.bold = True
                set_font(para, font_name='Arial')  # You can change to Calibri or Times New Roman

    return doc


@app.route('/', methods=['GET', 'POST'])
def process_text():
    matched_codes = []
    combined_text = ""
    prompts = [
            {
                    "prompt": "C reate an outline for the assignment. Provide a breakdown of sections with a suggested word count for each, totalling to required words as per the Assignment Brief . Ensure that the outline aligns with main focus, requirements and includes recommended headings or subheadings.Do not include word counts in the response. Do not include any disclaimers or assistance text in the response.",
                    "expected": "A detailed outline with sections, word allocations, and subheadings tailored to the topic."
            },
            {
                    "prompt": "Draft the detailed sections till the first 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief.Do not include word counts in the response. Do not include any disclaimers or assistance text in the response.",
                    "expected": "A well-rounded introduction with relevant background, 2-3 in-text citations, and properly formatted references."
            },
            {
                    "prompt": "Kindly provide the complete list of references used in the above 100 words in the reference format mentioned in the assignment brief.Do not include word counts in the response. Do not include any disclaimers or assistance text in the response.",
                    "expected": "A well-formatted list of references"
            },
            {
                    "prompt": "Draft the detailed next sections till the 2nd 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 2nd 1000 words in the reference format mentioned in the assignment brief.Do not include word counts in the response.",
                    "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
            },
            {
                    "prompt": "Draft the detailed next sections till the 3rd 100 words, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide The complete list of references used in the 3rd 1000 words in the reference format mentioned in the assignment brief.Do not include word counts in the response.",
                    "expected": "A well-rounded next section/s, 2-3 in-text citations, and properly formatted references and well-formatted list of references"
            },
            {
                    "prompt": "Draft the detailed next sections till the conclusion of the Assignment in the allocated word count according to the outline, these section should include 2-3 in-text references from reputable sources to establish the foundation of the topic per 200 words. Format references in the referencing style mentioned in the assignment brief also provide the complete list of references used in these sections in the reference format mentioned in the assignment brief.Do not include word counts in the response.",
                    "expected": "A well-rounded next section/s till Conclusion, 2-3 in-text citations, and properly formatted references and well-formatted list of references."
            }
        ]

    if request.method == 'POST':
        user_text = request.form.get('user_text', '').strip()
        if not user_text:
            flash('No text entered. Please provide input.')
            logging.warning("User submitted empty text input.")
            return render_template('input.html')

        # Check for matches in the user text
        found_codes = []
        for code in GPT_BOTS.keys():
            if re.search(rf'\b{re.escape(code)}\b', user_text, re.IGNORECASE):
                found_codes.append(code)

        if not found_codes:
            flash('No matching code found in the provided text.')
            logging.info("No matching codes found in user input.")
            return render_template('input.html')

        logging.info(f"Found matching codes: {found_codes}")

        all_responses = []
        # Process each matched code
        for code in found_codes:
            bot_url = GPT_BOTS[code]
            logging.info(f"Processing bot for code: {code} with URL: {bot_url}")

            driver = init_driver()
            try:
                bot_responses = process_bot(driver, bot_url, prompts)
                if bot_responses:
                    all_responses.extend(bot_responses)
                    logging.info(f"Collected {len(bot_responses)} responses from bot {code}.")
                else:
                    logging.warning(f"No responses collected from bot {code}.")
            except Exception as e:
                logging.error(f"Error processing bot for code {code}: {e}")
                flash(f"Error processing bot for code {code}: {e}")
            finally:
                try:
                    driver.quit()
                    logging.info(f"Closed Selenium WebDriver for bot {code}.")
                except Exception as e:
                    logging.error(f"Error closing WebDriver for bot {code}: {e}")

            matched_codes.append({"code": code, "bot_url": bot_url})

        if not all_responses:
            flash("No responses were generated from the bots.")
            logging.warning("No responses were collected from any bot.")
            return render_template('input.html', matched_codes=matched_codes)

        # Combine all responses
        for response in all_responses:
            combined_text += f"\n\n{response['answer']}"

        # Create a Word document
        try:
            doc = DocxDocument()
            doc.add_heading('Bot Responses', 0)
            doc.add_paragraph(combined_text)

            # Apply formatting to the document after creation
            doc = apply_formatting_to_word(doc)

            output_filename = "Bot_Responses_Formatted.docx"
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            doc.save(output_path)

            flash('Processing complete. Download the generated Word document below.')
            logging.info("Word document created and formatted successfully.")
            return render_template('input.html', 
                                   matched_codes=matched_codes,
                                   combined_text=combined_text,
                                   download_link=output_filename)
        except Exception as e:
            logging.error(f"Error creating Word document: {e}")
            flash(f"Error creating Word document: {e}")
            return render_template('input.html', matched_codes=matched_codes, combined_text=combined_text)

    # GET request
    return render_template('input.html')

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        logging.info(f"User requested download for file: {filename}")
        return send_file(file_path, as_attachment=True, download_name=filename)
    else:
        flash('File not found.')
        logging.warning(f"Download attempted for non-existent file: {filename}")
        return render_template('input.html')

if __name__ == '__main__':
    app.run(debug=True)

